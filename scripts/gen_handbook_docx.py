from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

BLUE = RGBColor(0x47, 0x85, 0xA8)
BLACK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)

def sf(run, name='宋体', size=11, bold=False, color=BLACK):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.element.rPr.rFonts.set(qn('w:eastAsia'), name)

def H(text, size=15):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    sf(r, '黑体', size, True, BLUE)
    return p

def P(text, size=11, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    sf(r, '宋体', size, bold)
    return p

def T(headers, rows):
    n = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t._tbl.tblPr.append(parse_xml(
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="AAAAAA"/>'
        '</w:tblBorders>'))
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=''
        p=c.paragraphs[0]; p.paragraph_format.line_spacing=1.2
        r=p.add_run(h); sf(r,'黑体',10,True,RGBColor(255,255,255))
        sh=parse_xml('<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="4785A8" w:val="clear"/>')
        c._tc.get_or_add_tcPr().append(sh)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; p.paragraph_format.line_spacing=1.2
            r=p.add_run(str(v)); sf(r,'宋体',9)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(6)
    return t

def Sep():
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(12)
    r=p.add_run('─'*55); sf(r,'宋体',8,color=GRAY)

# === COVER ===
for _ in range(4): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('🌊 出海云'); sf(r,'黑体',28,True,BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('产 品 推 广 手 册'); sf(r,'黑体',20,True,BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(16)
r=p.add_run('全产品线 · 给博主的一站式推广资料包'); sf(r,'宋体',13,color=GRAY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(6)
r=p.add_run('版本 v1.1 ｜ 2026年7月 ｜ 覆盖 8 款产品'); sf(r,'宋体',10,color=GRAY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(50)
r=p.add_run('⭐ GitHub 仓库：github.com/mars1417/mars-toolkit'); sf(r,'宋体',10,True,BLUE)
doc.add_page_break()

# === 1 ===
H('一、品牌简介')
P('出海云 是一个面向个人用户的技术工具产品系列，涵盖网络加速节点、系统自动化工具、运维保障工具三大品类。用户一次性付费购买后即可使用，无需月月续费。')
T(['项目','说明'],[
    ['品牌定位','人人买得起的实用技术工具'],
    ['产品线','订阅节点 + 一次性买断工具 + 免费工具'],
    ['更新频率','每周三/周五定期更新产品线'],
    ['目标用户','留学生、外贸从业者、技术爱好者、普通网民'],
    ['价格区间','¥0 ～ ¥99'],
    ['购买方式','官网直接下单，微信付款，全自动交付'],
])
P('品牌定位',bold=True)
P('「好用的技术工具，不一定很贵。¥9.9 起步，一次购买，长期使用。」')
P('品牌优势',bold=True)
T(['优势','说明'],[
    ['✅ 价格低','¥9.9 起步，远低于市面同类'],
    ['✅ 一次购买','买断制为主，不用月月续费'],
    ['✅ 自动运行','买了就能用，不需要技术背景'],
    ['✅ 售后兜底','3 天无理由退款，有问题人工处理'],
    ['✅ 持续更新','工具持续维护，不会买了就废'],
])
Sep()

# === 2 ===
H('二、全产品线总览')
P('订阅制产品（出海云系列）',bold=True)
T(['产品','价格','说明','状态'],[
    ['🌊 出海云·月付','¥9.9/月','50 条优选节点','✅ 已上线'],
    ['🌊 出海云·包年','¥150/年','送 Auto-Pilot（价值 ¥9.9）','✅ 已上线'],
])
P('')
P('一次性买断工具',bold=True)
T(['产品','价格','一句话说明','状态'],[
    ['🍎 macOS Care Pack','¥0 免费','Mac 系统优化工具集，免费下载','✅ 已上线'],
    ['🔄 VPN Auto-Pilot','¥9.9','V2Box 自动轮换工具','✅ 已上线'],
    ['🛡️ Tunnel Guardian','¥9.9','内网隧道自愈工具','✅ 已上线'],
    ['🌐 Cloud Ladder Bundle','¥29.9','Auto-Pilot + Guardian 合体包','✅ 已上线'],
])
P('')
P('预售产品',bold=True)
T(['产品','价格','预计上线','说明'],[
    ['💎 Pro Lite','¥49.9','7月底','更高阶的工具套装'],
    ['👑 Pro Max','¥99','8月','旗舰级全能工具包'],
])
P('')
P('推荐组合方案',bold=True)
T(['场景','推荐组合','总价'],[
    ['🥇 入门（仅上网）','出海云月付 ¥9.9','¥9.9'],
    ['🥇 标准（上网+自动）','出海云包年 + Auto-Pilot（送）','¥150'],
    ['🥇 全都要（上网+隧道+自动）','出海云包年 + Guardian + Care Pack','¥159.9'],
    ['🥇 技术宅全家桶（不上网也要买）','Auto-Pilot + Guardian + Care Pack','¥19.8'],
])
Sep()

# === 3 ===
H('三、产品详解')
P('3.1 🌊 出海云·月付 ¥9.9/月',bold=True)
P('一句话卖点：一杯奶茶钱，50 条优选节点，不限流量。')
P('适合谁：需要访问海外网络查资料、看 YouTube、用 ChatGPT 的用户；不想花大价钱买机场的；免费节点总失效的。')
P('推广重点：「¥9.9 一个月，一杯奶茶钱。50 条节点自动更新，比那些便宜的共享机场稳太多了。」')

P('')
P('3.2 🌊 出海云·包年 ¥150/年',bold=True)
P('一句话卖点：包年送工具，折合 ¥12.5/月，比月付省 37%。')
P('额外赠送：VPN Auto-Pilot（价值 ¥9.9）')

P('')
P('3.3 🍎 macOS Care Pack（免费 ¥0）',bold=True)
P('一句话卖点：Mac 优化工具包，免费下载，买了就赚。')
P('解决什么问题：Mac 用久了卡顿、磁盘空间不够。一键优化系统缓存、释放内存。')
P('适合谁：Mac 用户、感觉电脑变慢的人、不想装那些又大又贵的 Mac 清理软件的人。')
P('推广重点：「免费的东西谁不爱？关键是真的能用。零成本，不拿白不拿。」')

P('')
P('3.4 🔄 VPN Auto-Pilot ¥9.9（一次性买断）',bold=True)
P('一句话卖点：V2Box 的自动驾驶，¥9.9 一劳永逸。')
P('解决什么问题：V2Box 节点过期后需手动换订阅。Auto-Pilot 帮你自动轮换、自动连最快节点。')
P('适合谁：已用 V2Box 但懒得手动维护的、追求省事的用户。')
P('推广重点：「¥9.9 买一次，以后再也不用手动换节点了。自动选最快的、断了自动重连。」')

P('')
P('3.5 🛡️ Tunnel Guardian ¥9.9（一次性买断）',bold=True)
P('一句话卖点：内网穿透不再掉线，¥9.9 买个安心。')
P('解决什么问题：cpolar 隧道重启 URL 就变了，所有服务全断。Guardian 自动检测变化并修复。')
P('适合谁：用 cpolar 做内网穿透的技术用户、自己搭服务的站长。')
P('推广重点：「用过 cpolar 的都懂——服务跑得好好的，一重启隧道链接变了全崩。¥9.9 买个省心。」')

P('')
P('3.6 🌐 Cloud Ladder Pro Bundle ¥29.9',bold=True)
P('一句话卖点：两个工具打包买，省 ¥0.9，适合技术用户一次配齐。')
P('包含：VPN Auto-Pilot（¥9.9）+ Tunnel Guardian（¥9.9）')
Sep()

# === 4 ===
H('四、目标用户画像')
T(['画像','需求','推荐产品'],[
    ['🧑‍💻 留学生','需要稳定翻墙工具','出海云月付或包年'],
    ['👨‍💼 外贸/跨境电商','访问海外网站、WhatsApp','出海云月付或包年'],
    ['🧑‍🎓 学生党','查资料、用 ChatGPT，预算有限','出海云月付¥9.9'],
    ['👨‍🔧 技术宅','翻墙+内网穿透+Mac优化全都要','出海云+Auto-Pilot+Guardian'],
    ['👨‍👩‍👧 普通Mac用户','电脑卡顿想优化','macOS Care Pack（免费）'],
    ['👨‍💻 站长/自建服务','内网穿透老断连','Tunnel Guardian'],
    ['🧑‍💼 轻度技术用户','有V2Box但不想折腾','VPN Auto-Pilot'],
])
Sep()

# === 5 ===
H('五、推广话术包')
P('版本 A：简单直接（朋友圈）',bold=True)
P('🌊 出海云 · 工具全家桶\n🍎 macOS Care Pack ➔ 免费下载\n🌊 出海云节点 ➔ ¥9.9/月\n🔄 VPN Auto-Pilot ➔ ¥9.9\n🛡️ Tunnel Guardian ➔ ¥9.9\n👉 https://mars1417.github.io/mars-toolkit/')
P('')
P('版本 B：痛点触动（Mac用户）',bold=True)
P('Mac 越来越慢？别急着换电脑。\n🍎 macOS Care Pack 免费工具，一键清理、释放内存\n零成本，用了就知道爽\n👇 顺手还能看看别的实用工具\nhttps://mars1417.github.io/mars-toolkit/')
P('')
P('版本 C：痛点触动（翻墙用户）',bold=True)
P('还在用那些满大街的共享机场？慢、限速、跑路三件套？\n出海云 ¥9.9/月，优选线路 50 条节点自动更新\n专业节点服务，稳得很')
P('')
P('版本 D：对比型',bold=True)
P('❌ 免费节点：不稳定、老失效、不敢用\n❌ 大机场：年付几百、节点动不动被墙\n✅ 出海云：¥9.9/月、优选线路、50 节点、不限流量')
P('')
P('公众号标题参考：',bold=True)
P('①《从翻墙到 Mac 优化，这些 ¥9.9 的工具我用了半年》\n②《2026 年个人技术工具箱推荐：总价不到 ¥30》\n③《Mac 卡顿不用换电脑，一个免费工具就能救回来》\n④《内网穿透老掉线？这个 ¥9.9 的工具自动帮你修》')
P('')
P('Telegram 文案：',bold=True)
P('🔥 出海云 · 技术工具全家桶\n✅ 出海云节点 ¥9.9/月 → 50 条优选节点\n✅ VPN Auto-Pilot ¥9.9 → V2Box 自动驾驶\n✅ Tunnel Guardian ¥9.9 → 内网隧道永不掉线\n✅ macOS Care Pack ¥0 → Mac 优化免费下\n\n👉 https://mars1417.github.io/mars-toolkit/')
P('')
P('推广禁忌：',bold=True)
T(['❌ 不要写','✅ 建议写'],[
    ['翻墙、VPN、梯子','出海加速、科学上网、网络优化'],
    ['突破封锁、翻墙软件','访问海外网络、全球加速'],
    ['免费翻墙、破解','高速节点、优选线路'],
])
Sep()

# === 6 ===
H('六、购买流程')
P('用户下单路径：')
P('① 打开链接 → https://mars1417.github.io/mars-toolkit/')
P('② 选择产品 → 出海云点「立即订阅」，工具点对应产品链接')
P('③ 进入下单页 → 点「下一步：扫码付款」')
P('④ 微信扫码付款')
P('⑤ 填信息提交 → 微信昵称 + 邮箱 + 交易单号后4位')
P('⑥ 成功页显示订阅链接或下载信息 → 全程 3 分钟')
P('')
P('博主引导下单话术：',bold=True)
P('「点这个链接进去，选你要的产品。微信扫码付款，付完填个邮箱。出海云的话页面上直接显示订阅链接，复制到 V2Box 就能用。买断工具页面直接下载。全程 3 分钟。」')
Sep()

# === 7 ===
H('七、常见问题 FAQ')
P('Q1：我不会配置，买了能用吗？',bold=True)
P('出海云：下单成功页直接给链接，复制到 V2Box → 点「导入订阅」→ 粘贴 → 自动连接。买断工具：页面提供下载，一键运行。')
P('Q2：支持什么客户端？',bold=True)
P('出海云：V2Box、v2rayN、Clash Meta、Quantumult X、Shadowrocket、Stash 等。买断工具：macOS 平台。')
P('Q3：速度怎么样？',bold=True)
P('优选 IP 自动选最低延迟节点，日常 YouTube 4K 流畅，延迟 100-200ms。')
P('Q4：节点会失效吗？',bold=True)
P('系统自动维护节点池，始终保持可用。')
P('Q5：买了以后能退款吗？',bold=True)
P('所有产品支持 3 天内无理由退款。直接微信联系，秒退。')
P('Q6：包年送的 VPN Auto-Pilot 是什么？',bold=True)
P('V2Box 配套工具，自动更新订阅、自动选最快节点、断线自动重连。单买 ¥9.9。')
P('Q7：macOS Care Pack 真的免费吗？',bold=True)
P('是的，¥0。基础版永远免费。')
P('Q8：VPN Auto-Pilot 和 Tunnel Guardian 有什么区别？',bold=True)
T(['','VPN Auto-Pilot','Tunnel Guardian'],[
    ['管什么','V2Box 节点轮换','cpolar 隧道修复'],
    ['解决','节点过期无人换','隧道重启 URL 变'],
    ['适合','翻墙用户','内网穿透用户'],
])
P('Q9：Pro Lite 和 Pro Max 什么时候出？',bold=True)
P('Pro Lite 预计 7 月底，Pro Max 预计 8 月。关注官网获取最新动态。')
Sep()

# === 8 ===
H('八、竞品对比')
P('出海云 vs 市面常见方案',bold=True)
T(['对比项','出海云 ¥9.9/月','免费节点','主流机场 ¥20-40/月','自建 VPS'],[
    ['价格','✅ 低','✅ 免费','⚠️ 中等','❌ ¥30-100/月'],
    ['稳定性','✅ 高','❌ 常失效','✅ 中等','⚠️ 取决于VPS'],
    ['速度','✅ 快（优选）','❌ 慢','✅ 看线路','⚠️ 看配置'],
    ['节点数','50','不定','20-100','1-3'],
    ['自动更新','✅','❌','✅','❌'],
    ['不限流量','✅','✅','❌ 常限','✅'],
    ['配置难度','⭐（复制链接）','⭐⭐⭐','⭐⭐','⭐⭐⭐⭐'],
])
P('')
P('Auto-Pilot vs 手动维护',bold=True)
T(['对比项','Auto-Pilot ¥9.9','手动维护'],[
    ['价格','¥9.9 一次','免费'],
    ['时间成本','0，自动运行','每天花时间手动换'],
    ['断线恢复','自动重连','发现断线才手动连'],
    ['节点选择','自动选最快的','手动试，靠直觉'],
])
P('')
P('一句话总结：出海云 = 自建节点的品质 + 机场的便宜价格。比免费节点稳定，比大机场便宜一半。买断工具 ¥9.9 起步，一次购买终身省心。')
Sep()

# === 9 ===
H('九、联系方式 & 售后')
T(['项目','说明'],[
    ['产品官网','https://mars1417.github.io/mars-toolkit/'],
    ['GitHub 仓库','https://github.com/mars1417/mars-toolkit（欢迎⭐）'],
    ['售后支持','通过微信联系（下单时填的微信号）'],
    ['退款政策','3 天内无理由退款'],
    ['售后时效','工作日 2 小时内响应'],
])
P('')
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before=Pt(20)
r=p.add_run('推广合作说明')
sf(r,'黑体',13,True,BLUE)
P('博主推广使用自有链接 → 用户访问 GP 页面 → 下单付款到我们。订单全自动处理，博主无需介入售后。欢迎长期合作 💪')

P('')
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('📅 更新说明：本手册每周三/周五随产品更新同步修订。博主每次推广前建议检查最新版本。')
sf(r,'宋体',9,color=GRAY)

doc.save('/Users/mars/Desktop/_项目/mars-toolkit/出海云_产品推广手册.docx')
print('DONE')
